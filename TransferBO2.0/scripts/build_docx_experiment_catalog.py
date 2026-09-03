"""Build the experiment-catalog DOCX (docs/28_experiment_catalog.md -> docx).

Structure mirrors docs/28: overview table, per-stage sections with design /
scale / results / conclusion, and each relevant figure embedded right after
its experiment. Chinese font = Microsoft YaHei (eastAsia).

Usage:
    python scripts/build_docx_experiment_catalog.py
Output:
    docs/TransferBO2.0_experiment_catalog.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "results" / "figures"
OUT = ROOT / "docs" / "TransferBO2.0_experiment_catalog.docx"

CN_FONT = "Microsoft YaHei"
DARK = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x1F, 0x77, 0xB4)
GRAY = RGBColor(0x59, 0x59, 0x59)

# experiment label -> figure files (embedded after the section)
FIG_MAP = {
    "胺化 LOSO 主表": ["step1_amination_effects.png", "step1_amination_bsf.png"],
    "Suzuki LOSO 主表": ["step1_suzuki_effects.png", "step1_suzuki_bsf.png"],
    "Rep-A Morgan 底物": ["step1b_repA_morgan_substrate.png"],
    "Rep-B DFT 条件试点": ["step1b_repB_dft_pilot.png"],
    "M1 init vs 续跑": ["step2_m1_init_vs_post.png"],
    "M2 池化 vs 近邻": ["step2_m2_pool_vs_nearest.png"],
    "Suzuki shared-init/matched-post": ["p0_matched_init.png"],
    "胺化 matched-init（C1/C2/C3）": [],
    "源数门槛+清单稳定性": ["p1p2_source_robustness.png"],
    "borylation（主外部库）": ["p4_borylation_effects.png", "p4_borylation_per_target.png"],
    "HiTEA Suzuki（第二外部库）": ["p4_hitea_effects.png", "p4_hitea_per_target.png"],
    "排序保持+双通道": ["rank_preservation.png"],
    "清单聚合规则": ["strategy_list_rules.png"],
    "续跑事前规则（additive R²）": ["strategy_continuation_c1.png"],
    "探针门 G2": ["strategy_probe_gate.png"],
    "rank_median AUC 复核": ["rankmed_audit_compare.png"],
    "四臂 warm 续跑": ["continuation_arms.png"],
    "CHAOS 一维": ["chaos_validation.png"],
}


def set_font(run, size=10.5, bold=False, color=None, name=CN_FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def para(doc, text="", size=10.5, bold=False, color=None, space_after=4, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {0: 20, 1: 15, 2: 12.5, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=sizes.get(level, 11), bold=True,
             color=DARK if level <= 1 else None)
    return p


def add_table(doc, header, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        set_font(hdr[i].paragraphs[0].add_run(str(h)), size=font_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = hdr[i]._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "1F3A5F"})
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            set_font(cells[i].paragraphs[0].add_run(str(v)), size=font_size)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


def add_fig(doc, name, caption):
    pth = FIG / name
    if not pth.exists():
        para(doc, f"[缺图: {name}]", size=9, color=GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(pth), width=Inches(5.9))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    set_font(cap.add_run(caption), size=8.5, color=GRAY)


def main() -> int:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    # ---------- cover ----------
    heading(doc, "TransferBO2.0 实验全览", 0)
    para(doc, "2026-08-24 整理 · 按研究阶段登记全部实验（设计用途 / 规模 / 主结果 / 结论 / 可视化图）",
         size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    para(doc, "数字均为锁定值（FROZEN_CLAIMS / P4 摘要 / 策略研究摘要），与 results/paper_numbers/manifest.md 核对一致。"
              "图存放 results/figures/，生成脚本 scripts/make_experiment_figs.py。",
         size=9.5, color=GRAY, space_after=8)
    para(doc, "推断口径：target 级（seed 平均 → 靶级 → 配对 bootstrap 95% CI，B=5000）；主指标 AUC@20。",
         size=9.5, color=GRAY, space_after=12)

    # ---------- overview table ----------
    heading(doc, "0. 总览表", 1)
    overview = [
        ["阶段", "实验名", "目录", "规模"],
        ["Step1 效应", "胺化 LOSO 主表", "results/amination_v1_full", "450 (15×6×5)"],
        ["Step1 效应", "Suzuki LOSO 主表", "…/suzuki_v1_full_rt/suzuki_v1_full", "360 (12×6×5)"],
        ["Step1 补充", "胺化 topk 消融", "results/amination_topk_ablation", "525"],
        ["Step1 补充", "pair 试点（单源→靶）", "amination/suzuki_pair_v1_pilot", "126×2"],
        ["Step1b", "Rep-A Morgan 底物", "rep_A_morgan_sub_full ×2", "450+360"],
        ["Step1b", "Rep-B DFT 条件试点", "suzuki_rep_B_dft_cond_pilot", "36"],
        ["Step1b", "Rep-B both Morgan", "morgan_both_full ×2", "450+360"],
        ["Step2 机制", "M1 init vs 续跑", "results/step2_m1", "胺化+Suzuki"],
        ["Step2 机制", "M2 池化 vs 近邻", "results/step2_m2", "四库"],
        ["Step3 P0", "Suzuki shared-init/matched-post", "results/suzuki_p0_shared_init", "360"],
        ["Step3 P0", "胺化 matched-init (C1/C2/C3)", "results/amination_matched_init_audit", "150"],
        ["Step3 P1/P2", "源数门槛+清单稳定性", "results/p1p2_source_robustness", "离线"],
        ["P4 外部验证", "borylation（主外部库）", "results/p4_borylation", "990"],
        ["P4 外部验证", "HiTEA Suzuki（第二外部库）", "results/p4_hitea", "330（08-24 修复重跑）"],
        ["机制", "排序保持+双通道", "results/rank_preservation", "离线"],
        ["策略研究", "清单聚合规则", "results/strategy_list_rules", "四库"],
        ["策略研究", "续跑事前规则 (additive R²)", "results/strategy_continuation", "离线"],
        ["策略研究", "探针门 G2", "results/strategy_probe_gate", "离线回放"],
        ["策略研究", "rank_median AUC 复核", "results/rankmed_audit_compare", "355"],
        ["策略研究", "四臂 warm 续跑", "results/continuation_arms_compare", "230 (Suzuki 类)"],
        ["边界验证", "CHAOS 一维", "results/chaos_validation", "100"],
    ]
    add_table(doc, overview[0], overview[1:], widths=[1.0, 2.3, 2.3, 1.3])

    def sec(title, lines, figs, table=None, table_w=None):
        heading(doc, title, 1)
        for ln in lines:
            if ln.startswith("**") and "**" in ln[2:]:
                para(doc, ln, size=10.5, bold=True)
            else:
                para(doc, ln, size=10.5)
        if table:
            add_table(doc, table[0], table[1:], widths=table_w)
        for f in figs:
            add_fig(doc, f, f"图：{f}")

    # ---------- section 1 ----------
    sec("1. Step1 效应（主锁，08-20）",
        ["1.1 胺化（Pd C–N，15×260，450 jobs）——设计：LOSO（历史=其余 14 底物多源池化），6 策略×5 seeds，回答“历史能否加速新底物 BO”。主结果（锁定）：",
         "1.2 Suzuki（Pd C–C，12×308，360 jobs）——cold vs random −57.7（仅 4/12 靶赢）→ Q1 失败=冷启动 BO 不可靠（备注，非 topk 无效）；topk vs cold +149.9 [+38.8, +269.8] 排除 0；vs random +92.2 [0.0, +186.5]（可行但脆，CI 贴 0）。",
         "1.3 补充：topk 消融（k=5≡k=10，k=5 是设计单元）；pair 试点（单源 topk 弱于池化，1.0 pair 负效应的机制线索）。"],
        FIG_MAP["胺化 LOSO 主表"] + FIG_MAP["Suzuki LOSO 主表"],
        table=(["策略", "AUC@20", "vs cold [95% CI]", "frac>cold", "vs random", "frac>rand"],
               [["topk_warm", "1359.5", "+160.2 [+108.1, +211.6]", "0.87", "+268.0 [+216.3, +316.1]", "1.00"],
                ["nearest", "1316.2", "+117.0", "0.87", "+224.7", "1.00"],
                ["sim_weighted", "1218.6", "+19.3（null）", "0.73", "+127.1", "1.00"],
                ["safe_gate", "1210.8", "+11.5（null）", "0.67", "+119.3", "1.00"],
                ["cold", "1199.3", "—", "—", "+107.8 [+75.2, +145.4]", "1.00"],
                ["random", "1091.5", "—", "—", "—", "—"]]),
        table_w=[1.2, 0.9, 1.9, 0.9, 1.7, 0.8])

    # ---------- section 2 ----------
    sec("2. Step1b 表示轴稳健性（08-21）",
        ["2.1 Rep-A 底物 Morgan——只换底物指纹（hashed→Morgan r2），topk 效应不变（健全）；nearest 大幅变强但 M2 判定：默认仍是池化 topk。",
         "2.2 Rep-B 条件 DFT 试点（36 jobs）——topk 相对 OHE −157；cold≈random → 不升全量；条件默认 OHE。both-Morgan 同样不支持升级。"],
        FIG_MAP["Rep-A Morgan 底物"] + FIG_MAP["Rep-B DFT 条件试点"])

    # ---------- section 3 ----------
    sec("3. Step2 机制（08-22）",
        ["3.1 M1 双通道分解（carried=init 通道 / post_lift=续跑通道）：胺化 +278/−118、Suzuki +134/+16、borylation +186/−78、HiTEA +46/−20 → 价值位置库相关（Suzuki 的绝对 topk post +189 为四库最大）。",
         "3.2 M2 池化 vs 近邻：胺化 +160.2 > nearest +117.0；Suzuki +149.9 > +24.0（hashed）/Morgan 下翻盘——多源池化是稳健默认（08-24 复核：nearest 跨库不一致、无法事前识别）。",
         "3.3 Morgan 机制（M2-C）：Morgan 下 nearest 翻盘 = 换源（胺化 100%）+ init max 提高（胺化 62.6→68.6、Suzuki 68.1→78.7），而非全局 Spearman 更准（胺化 0.761→0.503 反降）——AUC 吃的是 init 最高点（M1），不是整体排序相关。"],
        FIG_MAP["M1 init vs 续跑"] + FIG_MAP["M2 池化 vs 近邻"] + ["step2_m2_morgan_mechanism.png"])

    # ---------- section 4 ----------
    sec("4. Step3 P0/P1/P2（08-22，策略研究前奏）",
        ["4.1 Suzuki shared-init（360 jobs）+ 胺化 matched-init（150 jobs）：Suzuki C1 ≈ +75（好起点后 EI 有价值）；胺化 C1 = +26.0（含 0）、C2 = +67.7（排除 0）→ “历史管起点，优化器管精修”。",
         "4.2 P1/P2 源数门槛：n=1 Jaccard≈0.17；≥3 启用、≥5 推荐；清单稳定性跨源差 → source coverage 强制上报。"],
        FIG_MAP["Suzuki shared-init/matched-post"] + FIG_MAP["源数门槛+清单稳定性"])

    # ---------- section 5 ----------
    sec("5. P4 外部验证（08-23；HiTEA 08-24 修复重跑）",
        ["5.1 borylation（Ni C–B，33×46，990 jobs，主外部库）：topk vs cold +107.6 [73.1, 144.9]、vs random +123.4 [89.1, 158.7]（均排除 0）；88%/94% 靶为正；init_best +8.63 排除 0、final_best +0.31≈0 → 强复现（init 模式），策略升“跨源已验证”。",
         "5.2 HiTEA Suzuki（Pd C–C，11×41–48，330 jobs，第二外部库，08-24 修复后重跑）：topk vs cold +26.3 [−32.2, +80.9]（方向正、CI 含 0）；final_best +0.22 含 0；C2 修复后 +20.4 排除 0 → 弱方向正（部分复现）；小空间+30% 失败压缩效应。",
         "⚠ 修复记录：原报 +47.0/+3.19 排除 0 系 ingest 因子列全 NULL → OHE 退化 → EI 顺序扫描伪影；08-24 修复（众数填充）+ 重跑 495 jobs，结论如实降级（docs/18 §8.3）。"],
        FIG_MAP["borylation（主外部库）"] + FIG_MAP["HiTEA Suzuki（第二外部库）"])

    # ---------- section 6 ----------
    sec("6. 排序保持机制（08-24）",
        ["设计：四库跨底物条件排序 Spearman + 池化 top-5 在靶内落位，回答“为什么排序可迁移、数值不可迁移”。",
         "结果：ρ = 胺化 0.577 / borylation 0.361 / Suzuki 0.264 / HiTEA 0.088 / CHAOS 0.694（五库全正）；顶部比整体更稳（22.7/260、14.6/46 vs 87.7/308、38/48）。",
         "化学解释：条件排序由配体/碱本征性质（模板通用）决定；产率水平由底物活性（特异）决定 → 排序可迁移、数值不可；清单带排序、GP 学数值；池化=跨底物排序投票。",
         "附：learnability 靶级 borylation +0.427 (p=0.01) 系哈希种子伪影，08-24 确定性修复后 +0.098 不显著（靶级显著证据撤回）。"],
        FIG_MAP["排序保持+双通道"])

    # ---------- section 7 ----------
    sec("7. 策略研究四组件（08-24）",
        ["7.1 清单规则：init 层 rank_median +1.78 显著；AUC@20 复核 pooled +1.5 持平、Suzuki +34.8（AUC@5 +20.3 显著）、HiTEA −30.0（不显著）→ 默认 mean；rank_median 仅 Suzuki 类可选；稀疏面板禁用。",
         "7.2 续跑规则：C1 库级差异（Suzuki +75.5 > 胺化 +26.0/borylation +13.6/HiTEA +20.5）→ init 型库 EI 可选、Suzuki 类 EI 必选；additive R² 分档事前规则不可行（p=0.69）。",
         "7.3 探针门 G2：round-1 的 5 点观测=探针（零额外成本）；探针有效性 +0.319（四库全正）、G2 选源池化 init_best 三库正；AUC@20 未验证 → 批量协议（湿实验）下待验证组件。",
         "7.4 rank_median AUC 复核（355 jobs）：pooled +1.5 [−14.8, +18.4]（持平）、Suzuki +34.8、HiTEA −30.0 → 不全面升级。",
         "7.5 四臂 warm 续跑（Suzuki 类 23 靶，230 jobs）：B vs A −59.1 [−139.0, −3.6] 显著负；C vs A −28.5 [−66.2, +4.3] 负趋势 → 历史数据不应以 warm 点进入 target GP；清单（init）是历史价值的正确载体。"],
        FIG_MAP["清单聚合规则"] + FIG_MAP["续跑事前规则（additive R²）"] + FIG_MAP["探针门 G2"]
        + FIG_MAP["rank_median AUC 复核"] + FIG_MAP["四臂 warm 续跑"])

    # ---------- section 8 ----------
    sec("8. CHAOS 一维边界验证（08-24）",
        ["设计：4 固定反应 × 720 共享添加剂（Science 2022）；条件空间一维；板内 z(log UV) 去水平保排序；检验“清单机制是否依赖多维条件结构”。",
         "结果：排序保持 0.694（五库最高）；topk 4/4 靶正（+8.1，n=4 方向性）；续跑零增益（topk vs topk_random Δ=0，清单一次吃光信号）→ 机制不依赖多维结构；一维下清单=全部信号（极端 init 模式）。"],
        FIG_MAP["CHAOS 一维"])

    # ---------- section 9 ----------
    heading(doc, "9. 汇总：证据链闭合图", 1)
    chain = ("Step1 效应（胺化 +160✅ / Suzuki +150 脆）→ Step1b 表示稳健（OHE+Morgan 收口）"
             "→ Step2 机制（M1 init/续跑分位、M2 池化>近邻、C1/C2 分离）"
             "→ P0 匹配初始化（C1 弱/C2 强）→ P1/P2 源数门槛（≥3/≥5、coverage）"
             "→ P4 外部验证（borylation 强复现 ✅ / HiTEA 弱方向正，08-24 修复重跑）"
             "→ 排序保持机制（五库 ρ 全正 + 顶部更稳 + 化学解释）"
             "→ 策略四组件（mean 清单 / 分库续跑 / G2 待验证 / coverage）"
             "→ 四臂 warm 负结果（历史只进 init）"
             "→ CHAOS 一维边界（机制不依赖多维结构）"
             "→ 正文 v2 + Claims register（docs/26）")
    para(doc, chain, size=10.5, space_after=8)

    # ---------- section 10 ----------
    sec("10. 复现说明",
        ["全部图：python scripts/make_experiment_figs.py（matplotlib，中文字体 Microsoft YaHei）",
         "论文数字：python scripts/make_paper_numbers_manifest.py → results/paper_numbers/manifest.md",
         "结果 JSON 路径见 docs/19_work_snapshot.md §9 与各实验 summary.md"],
        [])

    doc.save(str(OUT))
    print(f"saved {OUT}  ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
