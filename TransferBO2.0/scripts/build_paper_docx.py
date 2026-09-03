"""Build TransferBO2.0 abstract + outline as .docx (from docs/21 v2, 主干版).

  python scripts/build_paper_docx.py
Output:
  docs/TransferBO2.0_abstract_outline.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "TransferBO2.0_abstract_outline_v2_trunk.docx"

TITLE = "TransferBO 2.0 — 论文摘要与大纲（v2 成文主干版）"
DATE = "2026-08-24"
ONE_LINER = (
    "目标：找一个正增益的历史数据应用策略。切入点的教训：单源 pair 迁移效应为负，改为多源池化后转正。"
    "发现：池化 top-5 条件清单是四库中唯一稳健的正增益形态；进 GP、门控无效。化学解释：同一反应模板内，"
    "条件好坏排序由配体/碱的本征性质决定（跨底物通用），产率水平由底物自身活性决定（底物特异）——"
    "排序可迁移、数值不可迁移；清单携带排序，GP 学的是数值。"
)

TITLES = [
    "定稿： “Throw most of it away: historical HTE data accelerates Bayesian optimization for new substrates through a five-condition list, not a transfer model”",
]

ABSTRACT_EN = (
    "Historical HTE data should help Bayesian optimization (BO) for new substrates, but how it enters the loop "
    "matters: single-source transfer gave negative effects in our earlier pair-based setup, while pooling the "
    "history across substrates turns it positive. Across four HTE libraries spanning three reaction classes "
    "(71 substrate-defined tasks; 2,130 leave-one-substrate-out runs; one frozen protocol with dual controls), "
    "the only robustly positive strategy is a pooled top-five condition list used as round-one initialization "
    "(+160 and +108 AUC vs. cold start on the two full-grid libraries; 88–94% of targets improved); injecting "
    "historical labels into the surrogate is null or negative everywhere. The chemistry behind this is simple: "
    "within one reaction template, the ranking of conditions is set by ligand/base properties that act on every "
    "substrate alike, while the level of yields is set by each substrate's own reactivity—so the ranking "
    "transfers across substrates and the magnitudes do not. A list carries the ranking; a GP learns magnitudes. "
    "Deployment rules follow: pool ≥3 history substrates (≥5 recommended), report per-condition source "
    "coverage, and choose the continuation (EI) by library type. We conclude with an honest scope: the "
    "five-condition list is the positive-gain strategy in the libraries we tested; its boundary is set by "
    "substrate diversity and template generality."
)

ABSTRACT_ZH = (
    "历史 HTE 数据理应能帮新底物的 BO，但以什么方式进入回路是关键：我们早先的单源 pair 切入效应为负，改用"
    "跨底物多源池化后效应转正。在四个 HTE 库、三个反应类（71 个底物任务、2,130 个 LOSO run、一套冻结协议与"
    "双对照）上，唯一稳健的正增益策略是多源池化 top-5 条件清单作第 1 轮（两个完整网格库 +160/+108 AUC vs cold，"
    "88–94% 靶提升）；把历史标签注入代理模型则在所有库上 null 或负。背后的化学很简单：同一反应模板内，条件的"
    "好坏排序由配体/碱对模板所有底物都成立的本征性质决定，产率的水平由每个底物自身的反应活性决定——所以排序"
    "跨底物可迁移，数值不可迁移；清单携带排序，GP 学的是数值。部署规则：≥3 个历史底物池化（推荐 ≥5）、逐条件"
    "上报 source coverage、续跑（EI）按库类型选择。结论范围如实收窄：五条件清单是我们在测试库中找到的正增益"
    "策略，其边界由底物多样性与模板通用性决定。"
)

OUTLINE = [
    ("1. Introduction（两段式）", [
        "问题与目标：找一个正增益的历史数据应用方式（不是检验冷启动 BO 本身）",
        "切入点的教训：单源 pair（负）→ 多源池化（正）——“怎么用历史”比“用不用历史”更重要",
    ]),
    ("2. Results（四节，主证据各一个表）", [
        "2.1 什么有效：池化 top-5 清单——主证据：胺化 +160.2、borylation +107.6（CI 排除 0，88–94% 靶为正）；更快而非更高（init_best 优势大、final_best ≈ 0）",
        "2.2 什么无效：历史标签进 GP（sim_weighted/contextual 四库 null）；门控（无可用门）——负结果与正结果同框，“扔掉大部分历史”是证据结论而非直觉退化",
        "2.3 为什么有效（化学性）：① 模板通用性：配体（位阻/供电子性）与碱（碱性/溶解性）决定条件好坏，对模板内所有底物成立；② 底物特异性：芳基卤电子效应/位阻决定产率水平，不翻转排序——除非极端位阻底物×极端位阻配体冲突（“部分保持”的化学来源，也因此只取顶部排序）；③ 池化的化学意义：多底物排序投票 = 平均掉底物特异噪声，留下模板通用主效应；④ 数值不可迁移：产率绝对值含底物活性因子——只用排序、不用数值",
        "2.4 怎么用（应用规则）：k=5 清单 + ≥3/≥5 源门槛 + coverage 上报 + 分库续跑（init 型库 EI 可选、Suzuki 类 EI 必选）；表述口径（相对指标，不承诺绝对产率）",
    ]),
    ("3. Methods（协议与统计，一页纸）", [
        "四库与来源；LOSO、k=5、OHE、GP-EI、n_init=5/B=20、seeds 0–4、双对照（vs cold 与 vs random）、靶级 bootstrap",
        "matched-init 审计（C1–C4，分离起点效应与过程效应）",
        "可复现声明（2,130 jobs、脚本、数据）",
    ]),
    ("4. Discussion（三个问题，各一段）", [
        "为什么进 GP 无用（化学：GP 学数值，数值含底物特异水平，不可迁移）",
        "为什么门控学不会（排序保持度可事后测但元特征猜不中——门的方向是探针直接测量，未来工作）",
        "边界与局限（四库范围、回顾性、无湿实验前瞻、seed=5——如实收窄，符合事实即可）",
    ]),
    ("5. Conclusion", [
        "五条件清单是正增益策略；化学根基 = 排序可迁移/数值不可迁移；范围 = 测试库与模板通用性之内",
    ]),
    ("SI（备忘录区：不占正文，供查证）", [
        "S1 四库与协议细节；S2 matched-init 全表；S3 轮次指标；S4 源数门槛曲线",
        "S5 聚合规则消融；S6 排序保持与双通道机制验证（results/rank_preservation/）",
        "S7 策略研究 Phase 0（元特征不可跨库判别——备忘录）；S8 批次单样本提示；S9 种子敏感性（待做）",
    ]),
]

REDLINES = [
    "主线永远讲“找策略”：目标 → 切入点（pair 负/池化正）→ 发现 → 化学解释 → 应用",
    "主证据只挑最有信服力的（胺化 +160、borylation +108 两个数字开路），其余进 SI",
    "化学解释必须落地到“配体/碱通用性质 vs 底物特异活性”，不堆机制验证细节",
    "负结果只说两层（进 GP 无效、门控不可行），不展开曲折过程",
    "结论范围可以收窄，但必须符合事实（四库、回顾性、边界如实）",
]


def main() -> int:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(TITLE, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"版本：{DATE}　|　来源：docs/21_paper_framework.md（唯一权威版本，本 docx 为导出件）")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("0. 主干（一句话）", level=1)
    doc.add_paragraph(ONE_LINER)

    doc.add_heading("1. Title", level=1)
    for t in TITLES:
        doc.add_paragraph(t)

    doc.add_heading("2. Abstract（EN）", level=1)
    doc.add_paragraph(ABSTRACT_EN)
    doc.add_heading("2. Abstract（ZH，供内部）", level=1)
    doc.add_paragraph(ABSTRACT_ZH)

    doc.add_heading("3. 正文结构（主干版）", level=1)
    for section, items in OUTLINE:
        doc.add_heading(section, level=2)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("4. 叙事红线", level=1)
    for r in REDLINES:
        doc.add_paragraph(r, style="List Number")

    doc.save(OUT)
    print(f"[OK] {OUT}  ({OUT.stat().st_size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
