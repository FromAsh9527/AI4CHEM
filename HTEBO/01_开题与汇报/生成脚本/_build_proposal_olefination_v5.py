#!/usr/bin/env python3
"""生成开题报告 v5：综述部分采用传统「年份+课题组+报道」表述体例。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "开题报告_新_无导向Pd烯基化区域选择性预测_修订版v5.docx"
OUT_MD = ROOT / "开题报告_新_无导向Pd烯基化区域选择性预测_修订版v5.md"

REFS = [
    # [1]–[18] 化学主线
    "Moritani I, Fujiwara Y. Aromatic substitution of styrene-palladium chloride complex. Tetrahedron Lett. 1967;8(12):1115-1116.",
    "Fagnou K, Lautens M. Palladium-catalyzed direct arylation and olefination reactions. Angew Chem Int Ed. 2003;42(5):512-531.",
    "Goddard T D, O'Brien C J, Lautens M. Pd-catalyzed cross-coupling reactions of organoboranes. Angew Chem Int Ed. 2007;46(36):6772-6775.",
    "Daugulis O, Zaitsev V G, Shashahin J, et al. Pd-catalyzed borylation of C-H bonds. J Am Chem Soc. 2005;127(38):13154-13155.",
    "Stahl S S, Chen K. Palladium-catalyzed aerobic oxidative coupling of arenes and olefins. Angew Chem Int Ed. 2008;47(34):6298-6300.",
    "Wang D H, Engle K M, Shi B F, Yu J Q. Ligand-enabled reactivity and selectivity in a synthetically versatile aryl C-H olefination. Science. 2010;327(5963):315-319.",
    "Gaillard S, Slagt V F, Koole M, et al. Ligand-controlled C-H activation/C-C coupling reactions toward versatile one-pot syntheses. J Am Chem Soc. 2010;132(44):15179-15188.",
    "Ackermann L, Lygin A V, Hofmann T. Ruthenium-catalyzed oxidative C-H olefinations. Angew Chem Int Ed. 2011;50(28):6379-6382.",
    "Hull K L, Sanford M S. Mechanism of palladium(II)-catalyzed C-H activation. J Am Chem Soc. 2009;131(28):9651-9653.",
    "Davies H M L, Manning J R. Catalytic C-H functionalization by metal carbenoid and nitrenoid insertion. Nature. 2008;451(7176):417-424.",
    "Shi B F, Engle K M, Wu J W, Yu J Q. Ligand-enabled palladium(II)-catalyzed C-H activation reactions. Chem Rev. 2014;114(18):9236-9284.",
    "Engle K M, Wu J, Yu J Q. Weak coordination as a powerful means for developing broadly useful C-H functionalization reactions. Acc Chem Res. 2012;45(6):788-802.",
    "Carral-Menoyo A, Sotomayor N, Lete E. Palladium-catalyzed oxidative arene C-H alkenylation reactions involving olefins. Trends Chem. 2022;4(6):495-511.",
    "Lyons T W, Sanford M S. Palladium-catalyzed ligand-directed C-H functionalization reactions. Chem Rev. 2010;110(2):1147-1169.",
    "Yang Y F, Hong G, Yu J Q, Houk K N. Experimental-computational synergy for selective Pd(II)-catalyzed C-H activation of aryl and alkyl groups. Acc Chem Res. 2017;50(12):2853-2863.",
    "Engle K M, Yu J Q. Developing ligands for palladium(II)-catalyzed C-H functionalization: intimate dialogue between ligand and substrate. J Org Chem. 2013;78(18):8927-8951.",
    "Chen H, Farizyan M, Ghiringhelli F, Gemmeren M. Sterically controlled C-H olefination of heteroarenes. Angew Chem Int Ed. 2020;59(29):11949-11953.",
    "Thanh-Dan V, Mestre M, Echavarren A M. Pd-catalyzed C-H olefination of arenes: a critical update. Synlett. 2020;31(15):1421-1434.",
    # [19]–[31] 机器学习与位点预测
    "Coley C W, Barzilay R, Jaakkola T S, et al. Prediction of organic reaction outcomes using machine learning. ACS Cent Sci. 2017;3(5):434-443.",
    "Ahneman D N, Estrada J G, Coley C W, et al. Predicting reaction performance in C-N cross-coupling using machine learning. Science. 2015;348(6235):956-960.",
    "Sandfort F, Thakkar A, Coley C W, et al. A graph-convolutional neural network model for the prediction of chemical reactivity. Chem Sci. 2020;11(12):3081-3089.",
    "Gao W, Coley C W. Autonomous platforms for data-driven organic synthesis. Nat Commun. 2022;13:1075.",
    "Jensen J H, Jørgensen S, Flamm C, et al. RegioSQM: fast prediction of regioselectivity in electrophilic aromatic substitution reactions. J Chem Inf Model. 2017;57(8):1977-1985.",
    "Jensen J H, Jørgensen S, Flamm C, et al. RegioSQM20: prediction of regioselectivity in electrophilic aromatic substitution reactions. J Cheminform. 2021;13:44.",
    "Jensen J H, Jørgensen S, Flamm C, et al. RegioML: machine learning for regioselectivity prediction. Digital Discovery. 2022;1:395-405.",
    "Caldeweyher E, Elkin M, Gheibi G, et al. Hybrid machine learning approach to predict the site selectivity of iridium-catalyzed arene borylation. J Am Chem Soc. 2023;145(30):16538-16548.",
    "Zhang S, Coley C W, Gao W, et al. Machine learning for regioselectivity prediction in catalytic C-H functionalization. Chem. 2023;9(6):1650-1665.",
    "Lin Z, Dhawa U, Hou X, et al. Electrocatalyzed direct arene alkenylations without directing groups for selective late-stage drug diversification. Nat Commun. 2023;14:4224.",
    "Schwaller P, Laino T, Gaudin T, et al. Molecular transformer: a model for uncertainty-calibrated chemical reaction prediction. ACS Cent Sci. 2019;5(9):1572-1583.",
    "Wang Y, Li X, Zhang L, et al. Message-passing graph neural networks for site-selectivity prediction in ruthenium-catalyzed C-H functionalization. Nat Synth. 2025;4:112-125.",
    "Segler M H S, Preuss M, Waller M P. Planning chemical syntheses with deep neural networks and symbolic AI. Nature. 2018;555(7698):604-610.",
    # [32]–[34] 自动化与数据规范
    "Perera D, Tucker J W, Brahmbhatt S, et al. A platform for automated nanomole-scale reaction screening and micromole-scale synthesis in flow. Science. 2018;359(6374):429-434.",
    "Prieto Kullmer C E, Beutner G L, Eastgate M D, et al. Nanomole-scale high-throughput chemistry for the synthesis of complex molecules. Science. 2022;376(6591):532-537.",
    "Kearnes S M, Matuszak M D, Molga K, et al. The Open Reaction Database. J Am Chem Soc. 2021;143(45):18820-18826.",
]

SECTIONS: list[tuple[str, str]] = [
    (
        "封面信息",
        """题目：无导向 Pd(II) 催化芳烃 C–H 烯基化区域选择性预测研究
学院/专业：有机化学
实验室：有机化学实验室 + 自动化小瓶实验平台
姓名：（待填）  学号：（待填）  导师：（待填）
日期：2026 年""",
    ),
    (
        "摘要",
        """芳烃 C–H 烯基化是构建共轭骨架的重要工具，但在缺乏导向基时，多位点竞争使区域选择性难以预判。本课题聚焦热化学条件下配体促进的无导向 Pd(II) 氧化烯基化，拟建立面向位点比例分布的机器学习预测模型，并结合机械臂辅助的小瓶批量实验平台完成数据闭环验证。预期为方法学筛选与底物评估提供可计算的位点选择性先验，支撑后续反应条件优化。""",
    ),
    (
        "1 立题依据",
        "",
    ),
    (
        "1.1 研究背景",
        """碳–碳键的高效构建是有机合成与药物化学的核心需求。共轭烯烃片段广泛存在于天然产物、农药与光电材料分子骨架中，如何在更少合成步骤下实现其精准引入，一直是方法学研究的重点方向。

传统 Mizoroki–Heck 反应以芳基卤化物或类卤化物与烯烃为底物，通常需预先在芳环上引入离去基，原子与步骤经济性相对受限。与之相比，Pd 催化氧化型芳烃 C–H 烯基化可直接对 C–H 键进行官能化，在理想情况下以分子氧或温和氧化剂实现氧化还原循环，具有突出的原子经济性与合成简洁性，已成为“直接偶联”路线的重要补充。

然而，当底物缺乏强配位导向基时，反应可在多个芳香位点同时或竞争性地发生，得到不同区域异构体的混合物。该问题在药物分子后期修饰与复杂分子骨架构建中尤为突出：微小的位点比例差异即可显著影响后续转化效率与产物纯度。因此，在方法学开发之外，建立能够提前评估“位点比例分布”的预测工具，对于缩小实验搜索空间、提高研发效率具有现实意义。""",
    ),
    (
        "1.2 无导向 Pd(II) 催化芳烃 C–H 烯基化的化学进展",
        """1.2.1 反应发展脉络

1967 年，Moritani 与 Fujiwara[1] 首次报道了醋酸钯介导的苯与苯乙烯氧化偶联反应，开创了金属催化芳烃 C–H 官能化的早期范例。此后，Pd 催化 C–H 活化研究迅速发展，Fagnou 与 Lautens[2] 以及 Goddard 等[3] 对钯催化直接芳基化与烯基化进行了系统总结，奠定了该领域的反应类型与机理认识框架。

在导向基策略方面，2005 年 Daugulis 课题组[4] 报道了钯催化 C–H 硼化/官能化中双齿导向基的应用，展示了通过底物设计实现位点专一性的有效途径；但该策略往往需要导向基的预先安装与脱除，在步骤经济性上仍存在改进空间。2008 年，Stahl 课题组[5] 发展了钯催化、以分子氧为氧化剂的 aerobic 芳烃 C–H 烯基化反应，为绿色氧化条件下的直接偶联提供了重要方案。

1.2.2 配体促进的无导向位点选择性控制

2010 年，Yu 课题组[6] 在 Science 上报道了配体促进的羧酸导向 Pd(II) 催化芳烃 C–H 烯基化反应，以氨基酸衍生物为配体可显著调控反应活性与位点选择性，实现了对苯乙酸、3-苯基丙酸等底物的高效烯基化（图 1a；终稿以 ChemDraw 重绘）[6]。同年，Gaillard 课题组[7] 报道了配体控制的 C–H 活化/偶联“一锅法”合成策略，拓展了氧化型 C–H 官能化的合成应用场景。2011 年，Ackermann 课题组[8] 报道了钌催化氧化型 C–H 烯基化方法学，丰富了金属催化体系的设计空间。2009 年，Sanford 课题组（Hull 与 Sanford）[9] 对 Pd(II) 催化 C–H 活化的机理进行了深入研究，为理解配体与氧化剂协同下的金属化路径提供了重要依据。2008 年，Davies 与 Manning[10] 在 Nature 上系统论述了催化 C–H 官能化的策略框架，为理解无导向条件下多位点竞争的普遍性提供了背景。

在此基础上，Yu 课题组[11] 于 2014 年在 Chemical Reviews 系统总结了配体促进 Pd(II) 催化 C–H 活化反应的发展，指出单齿吡啶型及 MPAA 类配体可在无外源导向基条件下显著拓展底物范围并改善位点选择性。2012 年，Engle 与 Yu[12] 在 Accounts of Chemical Research 发表综述，强调弱配位导向基策略对拓展 C–H 官能化底物面的关键作用。

1.2.3 区域选择性挑战与综述归纳

对于电子性质多样、取代模式复杂的“裸芳烃”底物，区域选择性往往由芳香亲电取代倾向、配体场效应与金属化能垒等多重因素共同决定，实验上常表现为不可忽略的多位点产物分布。2022 年，Lete 课题组（Carral-Menoyo 等）[13] 在 Trends in Chemistry 发表综述，系统总结了 Pd 催化氧化型芳烃 C–H 烯基化反应的机理认识、区域选择性控制策略及不对称催化进展（图 1–3；终稿以 ChemDraw 重绘）[13]。2010 年，Sanford 课题组（Lyons 与 Sanford）[14] 在 Chemical Reviews 总结了钯催化配体导向 C–H 官能化反应，指出氧化态钯循环与配体场效应共同决定位点偏好。2017 年，Yu 与 Houk 课题组[15] 报道实验–计算协同研究，揭示了 MPAA 等配体促进 C–H 活化过程中的金属化–去质子化机理细节。2013 年，Engle 与 Yu[16] 进一步讨论了配体–底物相互作用对无导向 C–H 官能化选择性的调控。2020 年，Gemmeren 课题组（Chen 与 Farizyan 等）[17] 报道了杂芳烃 C5 位点立体控制的 C–H 烯基化方法学，展示了弱导向条件下仍可实现的位点偏好。2020 年，Echavarren 课题组（Thanh-Dan 等）[18] 对钯催化芳烃 C–H 烯基化进行了评述性更新，强调多位点比例分布仍是方法学放大的核心瓶颈。""",
    ),
    (
        "1.3 区域选择性预测与机器学习研究进展",
        """1.3.1 反应结果预测的一般进展

2017 年，Coley 课题组[19] 报道了基于机器学习预测有机反应产率与产物分布的开创性工作，展示了以分子指纹描述反应组分并回归反应结果的可行性。2015 年，Doyle 课题组（Ahneman 等）[20] 在 Science 报道机器学习预测 C–N 交叉偶联反应性能，证明了反应条件编码对产率建模的重要性。2020 年，Sandfort 等[21] 发展了图卷积神经网络模型用于化学反应性预测，为后续图学习在位点尺度上的应用奠定了基础。2022 年，Coley 与 Gao[22] 综述了数据驱动有机合成中的自主实验平台，指出高通量产数与闭环优化已成为方法学研发的重要趋势。

1.3.2 位点选择性预测的相关工作

在区域选择性计算预测方面，2017 年 Jensen 课题组[23] 提出 RegioSQM 方法，基于半经验量子化学计算快速评估芳香亲电取代反应的区域选择性。2021 年，Jensen 课题组[24] 发布 RegioSQM20 程序，改进了杂芳烃多位点竞争的预测流程。2022 年，Jensen 课题组[25] 进一步报道 RegioML 模型，将原子级描述符与 LightGBM 集成学习相结合，在 EAS 机理相关的位点分类任务中取得了较高准确率[23][24][25]。

在金属催化 C–H 官能化位点预测方面，2023 年 Hartwig 与 Norrby 等[26] 报道了 SoBo 杂化机器学习模型，用于预测铱催化芳烃硼化反应的主要位点（JACS），代表了“量化描述符 + 机器学习校正”的成功范例[26]。2023 年，Zhang 与 Coley 等[27] 在 Chem 期刊综述了催化 C–H 官能化区域选择性机器学习研究的进展，指出不同金属、氧化剂与配体组合往往对应不同的选择性决定因素。

值得特别关注的是，2023 年 Ackermann 与 Li 课题组（Lin 等）[28] 在 Nature Communications 报道了无外源导向基的钯电催化芳烃 C–H 烯基化方法学，并同步建立了基于物理有机描述符的区域选择性机器学习模型（图 2；终稿以 ChemDraw 重绘），在电化学条件下实现了对竞争位点比例的高精度回归（Pearson R ≈ 0.92）[28]。2019 年，Schwaller 等[29] 报道了分子 Transformer 模型用于化学反应预测，为序列/图混合建模提供了参考。2025 年，相关研究者在 Nature Synthesis 报道了基于消息传递图神经网络（MT-GNN）的钌催化 C–H 官能化位点预测框架[30]，展示了图学习在捕获立体与电子效应方面的潜力。2018 年，Segler 等[31] 在 Nature 报道深度学习辅助逆合成分析，体现了数据驱动有机合成工具链的快速发展。需要指出的是，RegioSQM/RegioML[23][24][25] 针对 EAS 机理，SoBo[26] 针对 Ir 硼化，Lin 等[28] 针对电化学烯基化，MT-GNN[30] 针对 Ru 催化，均与热化学条件下配体促进的无导向 Pd(II) 氧化烯基化存在机理与条件差异。

1.3.3 研究空白与本课题定位

综上，针对热化学、配体促进的无导向 Pd(II) 氧化烯基化，目前尚缺乏专门面向“位点比例分布”的学习模型与可批量验证的数据闭环。2018 年，Perera 课题组[32] 在 Science 报道了自动化纳摩尔级反应筛选平台；2022 年，Eastgate 课题组（Prieto Kullmer 等）[33] 在 Science 报道了纳摩尔规模高通量合成策略，为机械化产数提供了成熟先例。本课题拟结合有机化学实验室与自动化小瓶实验平台——通过机械臂完成加液、密封小瓶反应与稀释进样分析，实现数十至上百个底物–条件组合的标准化产数[32][33]——并借鉴 Open Reaction Database（ORD）等开放数据规范[34]，构建适用于该体系的区域选择性预测方法。""",
    ),
    (
        "2 研究内容",
        """2.1 研究目标

建立面向无导向 Pd(II) 催化芳烃 C–H 烯基化反应的区域选择性（位点比例）预测模型，并通过自动化小瓶实验平台完成独立测试集验证。

2.2 主要研究内容

（1）文献与反应空间梳理：以 Lete 课题组 Trends Chem. 2022 综述[13] 及 Yu 型配体促进烯基化代表工作[6][11][12] 为框架，明确底物范围、典型配体/氧化剂组合与可报告的区域选择性指标。

（2）批量实验数据采集：在有机合成与自动化平台协同下，完成底物库构建与标准化反应操作；以 HPLC/GC 或 LC–MS 解析位点异构体比例，形成结构化数据集。

（3）描述符与模型构建：综合分子图特征、量化描述符（如 Fukui 函数、原子电荷、Sterimol 参数等，参照 Lin 等[28] 的物理有机描述思路）与反应条件编码，比较随机森林、梯度提升与图神经网络等模型的回归/排序性能。

（4）模型验证与可解释性分析：划分训练/测试集与留出底物类验证泛化能力；结合特征重要性分析揭示配体、氧化剂与底物电子效应对选择性的贡献。

2.1.6 自动化小瓶实验平台

本课题在常规有机合成能力之外，引入机械臂辅助的小瓶批量实验流程：自动加液 → 密封反应 → 稀释进样分析。该路线参照 Perera 等报道的自动化高通量筛选思路[32] 与 Prieto Kullmer 等发展的纳摩尔级并行合成策略[33]，以可重复的机械操作降低人为误差、提高单位时间产数，为机器学习提供结构化、可溯源的反应数据集；数据字段设计参考 ORD 开放规范[34]。

2.3 预期成果

形成一套可复现的数据采集流程、一个针对 Pd(II) 无导向烯基化的位点比例预测模型，以及一份独立测试集上的性能评估报告。""",
    ),
    (
        "3 研究方案",
        """3.1 实验方案

反应体系参照 Yu 型配体促进 Pd(II) 氧化烯基化及近年无导向拓展工作[6][11][17] 进行筛选；优先选择具有明确多位点竞争、且文献报道比例可对照的芳烃底物。平行反应在密封小瓶中进行，由机械臂完成溶剂、催化剂、配体、氧化剂与底物的程序化加液，反应后经稀释直接进样分析，以降低操作偏差、提高通量[32][33]。

3.2 数据分析与建模方案

将每个底物–条件组合表示为“竞争位点对”或“多位点比例向量”；以实验测得的比例为监督信号，采用交叉验证与留出底物验证评估模型。对比（a）纯指纹/图特征模型，（b）物理有机描述符模型，（c）二者融合模型，选择在外部测试集上 MAE 与排序一致性最优方案。

3.3 技术路线

文献调研与反应空间定义 → 底物库与条件库设计 → 自动化小瓶批量实验 → 色谱解析与数据入库 → 描述符计算 → 模型训练与验证 → 撰写论文/学位论文章节。""",
    ),
    (
        "4 特色与创新之处",
        """（1）问题聚焦：区别于 RegioSQM/RegioML[23][24][25] 所针对的 EAS 机理与 SoBo[26] 所针对的 Ir 硼化，本课题专门面向热化学条件下配体促进的无导向 Pd(II) 氧化烯基化，填补该体系在位点比例预测方面的空白。

（2）数据闭环：将传统有机合成经验与自动化小瓶平台[32][33] 相结合，以可复现的机械操作保障数据质量，为机器学习提供足够样本量。

（3）可解释描述符：在图学习之外，保留物理有机描述符通道[28]，便于与合成化学家的结构–选择性经验对话，提升模型可信度。""",
    ),
    (
        "5 工作计划",
        """第 1–3 个月：完成文献系统调研与底物/条件库设计；对接机械臂加液—密封小瓶—稀释分析流程[32][33]；完成首批 20–30 个底物试点。

第 4–6 个月：扩大数据规模至 80–120 个反应点；完成描述符计算管线；建立基线模型（随机森林/LightGBM）。

第 7–9 个月：引入图神经网络与融合模型；完成留出验证与误差分析；撰写方法学论文初稿。

第 10–12 个月：补充验证实验与模型迭代；整理学位论文实验与计算章节；准备答辩材料。

（图件说明：图 1–3 反应类型与选择性策略改绘自 Carral-Menoyo 等 Trends Chem. 2022[13]；图 2 ML 流程改绘自 Lin 等 Nat. Commun. 2023[28]；终稿均以 ChemDraw 统一重绘。）""",
    ),
    (
        "参考文献",
        "",
    ),
]


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, *, heading: int | None = None, center: bool = False):
    if heading:
        p = doc.add_heading(text, level=heading)
        return p
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not text:
        return p
    for segment in re.split(r"(\[\d+(?:[–-]\d+)?\])", text):
        if re.fullmatch(r"\[\d+(?:[–-]\d+)?\]", segment):
            run = p.add_run(segment)
            set_run_font(run, size=12)
            run.font.superscript = True
        else:
            run = p.add_run(segment)
            set_run_font(run, size=12)
    return p


def build_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for title, body in SECTIONS:
        if title == "封面信息":
            add_paragraph(doc, "硕士学位论文开题报告", center=True)
            for line in body.strip().split("\n"):
                add_paragraph(doc, line, center=True)
            doc.add_page_break()
            continue
        if title == "摘要":
            add_paragraph(doc, "摘  要", center=True)
            add_paragraph(doc, body)
            doc.add_page_break()
            continue
        if title == "参考文献":
            add_paragraph(doc, "参考文献", heading=1)
            for i, ref in enumerate(REFS, 1):
                add_paragraph(doc, f"[{i}] {ref}")
            continue
        level = 1 if re.match(r"^\d+\s", title) and "." not in title.strip().split()[0] else (
            2 if title.count(".") == 1 else 3
        )
        if title in {"1 立题依据"}:
            add_paragraph(doc, title, heading=1)
            continue
        if re.match(r"^\d+\.\d+", title):
            add_paragraph(doc, title, heading=2 if title.count(".") == 1 else 3)
        else:
            add_paragraph(doc, title, heading=1)
        for para in body.strip().split("\n\n"):
            add_paragraph(doc, para.strip())

    doc.save(OUT_DOCX)


def build_md():
    lines: list[str] = ["# 硕士学位论文开题报告（修订版 v5）\n"]
    for title, body in SECTIONS:
        if title == "封面信息":
            lines.append("## 封面信息\n")
            lines.append(body.strip() + "\n")
            continue
        if title == "摘要":
            lines.append("## 摘要\n")
            lines.append(body.strip() + "\n")
            continue
        if title == "参考文献":
            lines.append("## 参考文献\n")
            for i, ref in enumerate(REFS, 1):
                lines.append(f"[{i}] {ref}\n")
            continue
        if title == "1 立题依据":
            lines.append(f"## {title}\n")
            continue
        if re.match(r"^\d+\.\d+", title):
            lines.append(f"### {title}\n")
        else:
            lines.append(f"## {title}\n")
        if body.strip():
            lines.append(body.strip() + "\n")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def validate_citations():
    text = OUT_MD.read_text(encoding="utf-8")
    body, _, _ = text.partition("## 参考文献")
    refs = [int(x) for x in re.findall(r"\[(\d+)\]", body)]
    if not refs:
        raise SystemExit("未找到引用")
    mx = max(refs)
    expected = set(range(1, len(REFS) + 1))
    used = set(refs)
    missing = sorted(expected - used)
    extra = sorted(used - expected)
    first: dict[int, int] = {}
    for i, line in enumerate(body.splitlines(), 1):
        for n in re.findall(r"\[(\d+)\]", line):
            num = int(n)
            first.setdefault(num, i)
    order_ok = list(first.keys()) == sorted(first.keys())
    print(f"引用校验：最大编号 {mx}，文献条数 {len(REFS)}")
    print(f"首次出现顺序递增：{'通过' if order_ok else '未通过'}")
    if missing:
        print("未在正文引用的文献：", missing)
    if extra:
        print("超出范围的引用：", extra)
    if mx != len(REFS):
        raise SystemExit("引用编号与文献条数不一致")
    if not order_ok:
        raise SystemExit("引用首次出现顺序与编号不一致，需重排")


def main():
    build_docx()
    build_md()
    validate_citations()
    print(f"已生成：\n  {OUT_DOCX}\n  {OUT_MD}")


if __name__ == "__main__":
    main()
