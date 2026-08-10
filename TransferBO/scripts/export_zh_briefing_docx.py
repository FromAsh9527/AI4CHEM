#!/usr/bin/env python
"""Export Chinese briefing Word doc for TransferBO / EDBO Suzuki main line."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIGS = ROOT / "docs" / "figs"
OUT = ROOT / "docs" / "briefings" / "briefing_zh_EDBO_Suzuki_v0.6.docx"
OUT_CN = ROOT / "docs" / "briefings" / "本工作说明_EDBO_Suzuki_v0.6.docx"


def set_run_font(run, *, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size_pt: float = 12):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_asia)


def add_heading_cn(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, east_asia="黑体", ascii_font="Arial", size_pt=16 if level == 1 else 14)
    return p


def add_para(doc: Document, text: str, *, first_line_indent: bool = True, bold: bool = False, size: float = 12):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, size_pt=size)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size_pt=12)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        set_run_font(run, east_asia="黑体", size_pt=10.5)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size_pt=10.5)
    doc.add_paragraph()


def try_add_figure(doc: Document, path: Path, caption: str, width_in: float = 5.8):
    if not path.exists():
        add_para(doc, f"［图文件缺失：{path.name}］", first_line_indent=False, bold=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, east_asia="楷体", size_pt=10.5)
    run.italic = True


def main() -> int:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("本工作说明")
    r.bold = True
    set_run_font(r, east_asia="黑体", ascii_font="Arial", size_pt=18)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(
        "共享反应条件库上的历史标签池化：\n"
        "为何在 EDBO Suzuki 贝叶斯优化中不是安全默认策略"
    )
    r.bold = True
    set_run_font(r, east_asia="黑体", ascii_font="Arial", size_pt=14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "项目：TransferBO / TransferGate 线 · 主证据：EDBO2021 Suzuki\n"
        "文稿对应：manuscript_draft_DD_v0.6 · 日期：2026-08-07\n"
        "用途：向合作者 / 导师说明科学问题、证据链与主张边界（中文简报）"
    )
    set_run_font(r, east_asia="宋体", size_pt=10.5)
    for run in meta.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    add_heading_cn(doc, "一、一句话结论", 1)
    add_para(
        doc,
        "在 EDBO Suzuki「共享多维反应条件库、跨底物」设定下，把历史底物的产率标签持续并入"
        "无任务身份（no task ID）的目标板贝叶斯优化（BO）代理模型，相对「只用目标板数据的冷启动」"
        "并不是安全默认策略：在中等实验预算（约 30–50 次目标板查询）上，平均表现为负迁移；"
        "该结论在匹配目标板初始化后仍成立；任务内秩变换与中等强度的源数据降权均未能稳定翻正。",
        bold=False,
    )
    add_para(
        doc,
        "更短的公式化表述：共享候选条件空间 X，并不意味着可以把跨底物历史产率 y 当作目标响应函数的可交换观测。",
        first_line_indent=True,
    )

    add_heading_cn(doc, "二、科学问题（我们到底在回答什么）", 1)
    add_para(
        doc,
        "高通量实验与闭环优化常在同一套离散条件库上更换底物。一块板做完后，历史产率要不要直接喂给下一块板的 BO？",
    )
    add_para(doc, "本稿锁定的问题不是「迁移学习永远无效」，而是更可操作的一句：", first_line_indent=True)
    add_para(
        doc,
        "相对目标板冷启动，无任务身份、持续并入历史标签的池化，是不是安全默认？",
        bold=True,
        first_line_indent=True,
    )
    add_para(
        doc,
        "主证据库：EDBO2021 Suzuki（Shields 等）。条件库 X = 配体 × 碱 × 溶剂；任务 = 不同底物（底物不进入 X）。"
        "对照边界：Doyle（另一化学体系 + OHE）、PK2022/CHAOS 一维添加剂库——只说明「别处可以正迁移」，"
        "不与 EDBO 做效应量横比。",
    )

    add_heading_cn(doc, "三、方法与协议（读结果前需知道的设定）", 1)
    add_bullet(doc, "代理模型：sklearn 高斯过程 + EI；目标板 n_init=20，总预算 B=100（init 计入 B）。")
    add_bullet(doc, "历史源标签：可持续进入 GP，但不占用目标板实验次数；源点上限约 m=150。")
    add_bullet(doc, "无 task ID：源与目标观测被当作同一潜在函数上的点。")
    add_bullet(doc, "表示：Morgan / DRFP / 条件 DFT（C1 三表示同向）。")
    add_bullet(doc, "推断单位：有向源→目标 pair（N≈56）；种子估计算法波动，不能把上千条轨迹当 IID。")
    add_bullet(doc, "主报告窗口：B=30–50；B=100 因冷启动接近天花板，只作后期对照。")

    add_heading_cn(doc, "四、方法阶梯（A0–A3）", 1)
    add_table(
        doc,
        ["臂", "历史数据怎么用", "回答的假设"],
        [
            ["A0 冷启动", "不用", "目标板基线"],
            ["A1 原始池化", "把源产率与目标观测直接混进同一 GP", "标签是否可交换？"],
            ["A2 秩/百分位池化", "任务内先做秩变换再池化", "是否只是产率标度问题？"],
            ["A3 源降权", "抬高源点观测噪声 α_src=1e-4/w_s，w∈{0.1,0.25,0.5}", "弱权重是否够用？"],
        ],
    )
    add_para(
        doc,
        "S0：匹配目标板初始化。主网格曾存在冷启动与标签臂 init 不一致的实现问题；S0 修复后 init 100% 对齐。"
        "A2/A3 均在 S0 路径上相对 S0 冷启动比较。",
    )

    add_heading_cn(doc, "五、主要结果（数字摘要）", 1)

    add_heading_cn(doc, "5.1 C1：三表示中段平均略负", 2)
    add_para(doc, "主网格上，原始标签池化相对冷启动的 pair 平均 Δfrac（B=40）：")
    add_table(
        doc,
        ["表示", "B=40 Δfrac", "约 95% CI（pair）", "B=100"],
        [
            ["Morgan", "−0.041", "[−0.057, −0.027]", "−0.010"],
            ["DRFP", "−0.036", "[−0.052, −0.022]", "−0.013"],
            ["DFT", "−0.036", "[−0.051, −0.023]", "−0.018"],
        ],
    )
    add_para(
        doc,
        "三表示高度同向，说明不是某一种指纹的特例。B=100 差异变小，主要因为冷启动已接近各目标最优（天花板/余量问题），"
        "故主文强调中段预算。",
    )
    try_add_figure(
        doc,
        FIGS / "fig_edbo_suzuki_C1_pair_delta_by_budget.png",
        "图 2　C1：三表示 pair 级 Δfrac 随预算变化（相对冷启动）",
    )

    add_heading_cn(doc, "5.2 S0：匹配初始化不翻号", 2)
    add_table(
        doc,
        ["表示", "主网格 B=40", "S0 匹配 init B=40"],
        [
            ["Morgan", "−0.041", "−0.034"],
            ["DFT", "−0.036", "−0.030"],
        ],
    )
    add_para(doc, "匹配 init 使负效应略减轻，但中段仍显著为负。A2/A3 以 S0 冷启动为共同对照。")
    try_add_figure(
        doc,
        FIGS / "fig_edbo_suzuki_s0_vs_main_pair_delta.png",
        "图 3　S0 匹配初始化 vs 主网格（稳健性）",
    )

    add_heading_cn(doc, "5.3 A2 / A3：简单补救不够", 2)
    add_para(doc, "相对 S0 冷启动，B=40 的 pair 平均 Δfrac：")
    add_table(
        doc,
        ["方法", "Morgan", "DFT"],
        [
            ["A1 原始池化", "−0.034", "−0.030"],
            ["A2 秩池化", "−0.026", "−0.029"],
            ["A3 w=0.25（代表档）", "≈ −0.036", "≈ −0.031"],
        ],
    )
    add_para(
        doc,
        "A2 略减负、未翻正 → 不完全是「绝对产率标度」问题。"
        "A3 三档权重几乎重合且 ≈ A1 → 在测试的实际降权范围内，简单抬高源观测噪声未能稳定扭转负迁移。",
    )
    add_para(
        doc,
        "A3 健全性检查（PASS）：权重确实进入 GP 并改变后验；但 w=0.1–0.5 时相对 w=1 后验变化很小，"
        "约 75% 的轨迹采集路径完全相同；极端 w=1e-4 才明显靠近冷启动（仅作实现对照，不作正式方法臂）。"
        "因此 A3 是「有解释的负结果 / 敏感性不足」，不是代码空转。",
    )
    try_add_figure(
        doc,
        FIGS / "fig_edbo_suzuki_ladder_A1A2A3_B40.png",
        "图 4　方法阶梯 A1–A3（B=40，相对 S0 冷启动；点为 pair）",
    )

    add_heading_cn(doc, "六、主张边界（什么能说 / 什么不能说）", 1)
    add_para(doc, "可以说：", bold=True, first_line_indent=False)
    add_bullet(doc, "在 EDBO Suzuki 共享条件库基准上，无 task ID 的持续历史标签池化不是安全默认。")
    add_bullet(doc, "结果与「任务不匹配（task mismatch）」相一致，但不断言这是唯一机制。")
    add_bullet(doc, "Doyle / PK 表明别处可以出现有利迁移——Suzuki 结果不是「禁止一切迁移」。")
    add_para(doc, "不可以说：", bold=True, first_line_indent=False)
    add_bullet(doc, "历史反应数据没用；一切迁移学习都无效。")
    add_bullet(doc, "带 task ID / 多任务 GP 也一定无效（那是下一篇工作）。")
    add_bullet(doc, "所有反应家族都负迁移（amination 等尚未作为本稿必证）。")
    add_bullet(doc, "任意权重下源降权都无效（只测了实际档 w=0.1–0.5）。")
    add_bullet(doc, "把 Doyle 与 EDBO 的效应量直接排名或合并平均。")

    add_heading_cn(doc, "七、与旧叙事的关系", 1)
    add_para(
        doc,
        "早期草稿（v0.5 及以前）曾以 CHAOS/Doyle 上「标签池化平均有帮助」为主叙事。"
        "在 EDBO Suzuki 完整网格与方法阶梯完成后，主结论已切换为："
        "共享多维条件库上的 task-agnostic 池化需要警惕负迁移，不能当作默认暖启动。"
        "英文稿见 docs/manuscript_draft_DD_v0.6.md；本文件为中文说明版，便于组内汇报与讨论。",
    )

    add_heading_cn(doc, "八、后续工作（本稿不做）", 1)
    add_bullet(doc, "S3：显式任务身份特征；S4：多任务 GP / ICM；S5：仅用历史数据指导初始化。")
    add_bullet(doc, "问题升级为「如何安全迁移」，宜独立成篇，不为本稿硬找正结果。")
    add_bullet(doc, "可选、不阻塞投稿：EDBO amination 最小验证（Morgan+DFT × 冷启动/原始池化 × 匹配 init）。")
    add_bullet(doc, "不做：xTB 全量、Doyle 多表示重跑、amination 全套 A2/A3（除非最小验证值得深入）。")

    add_heading_cn(doc, "九、关键路径索引", 1)
    add_bullet(doc, "英文主稿：docs/manuscript_draft_DD_v0.6.md")
    add_bullet(doc, "主张冻结：results/paper_stats/FROZEN_CLAIMS.md")
    add_bullet(doc, "数字唯一源：results/paper_stats/EXPERIMENT_SUMMARY.md")
    add_bullet(doc, "成稿计划：docs/行动方案_成稿收尾.md")
    add_bullet(doc, "结果目录：results/external_edbo_suzuki{,_s0,_a2,_a3}/")
    add_bullet(doc, "主图：docs/figs/fig_edbo_suzuki_* ；示意图 fig1_same_library_transfer_schematic.png")

    add_heading_cn(doc, "十、给非本方向读者的读法", 1)
    add_para(
        doc,
        "若只看一张图：看图 4——三种「把历史产率塞进同一个无任务身份 GP」的用法，"
        "在中段预算上相对冷启动平均都不占优。"
        "若只记一句话：共用条件库 ≠ 可以安全地把旧底物产率当成新底物的标签。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    # also save under Chinese filename when filesystem encoding allows
    try:
        doc.save(OUT_CN)
        print("wrote", OUT_CN)
    except OSError as e:
        print("skip CN filename:", e)
    print("wrote", OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
