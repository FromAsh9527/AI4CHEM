#!/usr/bin/env python
"""Export manuscript_draft_DD.md to a Word .docx with embedded figures."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "manuscript_draft_DD.md"
FIGS = ROOT / "docs" / "figs"
OUT = ROOT / "exports" / "manuscript_draft_DD.docx"


def set_run_font(run, name: str = "Times New Roman", size: int = 11, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc: Document, text: str, *, style: str | None = None, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    # light markdown: **bold**, *italic*, `code`, ^{note}
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\\\([^)]+\\\))", text)
    if parts == [text]:
        # also split simple ** **
        parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True, size=size)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            set_run_font(run, italic=True, size=size)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=size - 1)
        else:
            # strip footnote markers like ^[...]
            clean = re.sub(r"\^\[[^\]]*\]", "", part)
            clean = clean.replace("\\(", "").replace("\\)", "")
            run = p.add_run(clean)
            set_run_font(run, bold=bold, italic=italic, size=size)
    return p


def add_image(doc: Document, rel_path: str, caption: str | None = None, width_in: float = 6.2):
    # rel_path like figs/xxx.png relative to docs/
    name = Path(rel_path).name
    path = FIGS / name
    if not path.exists():
        path = ROOT / "exports" / "paper_figs" / name
    if not path.exists():
        add_para(doc, f"[Missing figure: {rel_path}]", italic=True, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, italic=True)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.text = row[j] if j < len(row) else ""
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=9, bold=(i == 0))
    doc.add_paragraph()


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def export() -> Path:
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    i = 0
    pending_caption = None
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # skip status / coauthor meta blocks at end lightly — keep content
        if stripped.startswith("![") and "](" in stripped:
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
            if m:
                alt, rel = m.group(1), m.group(2)
                # look ahead for italic caption line
                cap = None
                if i + 1 < len(lines) and lines[i + 1].strip().startswith("*Fig"):
                    cap = lines[i + 1].strip().strip("*")
                    i += 1
                add_image(doc, rel, caption=cap or alt)
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            for run in p.runs:
                set_run_font(run, size=16, bold=True)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:].strip(), level=1)
            for run in p.runs:
                set_run_font(run, size=14, bold=True)
            i += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            rows, i = parse_md_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("$$") or stripped == "\\[":
            # gather math block as monospace paragraph
            buf = [stripped]
            i += 1
            while i < len(lines) and not (lines[i].strip().startswith("$$") or lines[i].strip() == "\\]"):
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                buf.append(lines[i])
                i += 1
            add_para(doc, " ".join(b.strip() for b in buf), italic=True, size=10)
            continue

        # display math starting with \[
        if stripped.startswith("\\["):
            buf = [stripped]
            i += 1
            while i < len(lines) and "\\]" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < len(lines):
                buf.append(lines[i])
                i += 1
            add_para(doc, " ".join(b.strip() for b in buf), italic=True, size=10)
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---"):
            i += 1
            continue

        if stripped.startswith("*Fig.") or stripped.startswith("*Fig "):
            # caption already handled with image; if orphan, add
            add_para(doc, stripped.strip("*"), italic=True, size=9)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2 and len(stripped) < 80:
            add_para(doc, stripped, bold=True, size=11)
            i += 1
            continue

        # bullet
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*.+?\*\*)", stripped[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    set_run_font(r, bold=True)
                else:
                    r = p.add_run(part)
                    set_run_font(r)
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            body = re.sub(r"^\d+\.\s", "", stripped)
            parts = re.split(r"(\*\*.+?\*\*)", body)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    set_run_font(r, bold=True)
                else:
                    r = p.add_run(part)
                    set_run_font(r)
            i += 1
            continue

        add_para(doc, stripped)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    alt = ROOT / "docs" / "manuscript_draft_DD.docx"
    try:
        doc.save(str(alt))
        print(f"Wrote {alt}")
    except PermissionError:
        print(f"Skipped locked {alt} (close Word and re-run if needed)")
    return OUT


if __name__ == "__main__":
    export()
